# -*- coding: utf-8 -*-
"""
Created on Tue Oct 28 07:00:54 2025

@author: vzocc
"""


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_global_style(docx_file):
    doc = Document(docx_file)

    # --- Helper for alignment names ---
    ALIGNMENT_NAMES = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }

    def get_paragraph_style(style_name):
        """Extract paragraph format and font info for a named style."""
        s = doc.styles[style_name]
        pf = s.paragraph_format
        font = s.font
        return {
            "font_name": font.name,
            "font_size": font.size.pt if font.size else None,
            "alignment": ALIGNMENT_NAMES.get(pf.alignment, "left"),
            "line_spacing": pf.line_spacing if pf.line_spacing else 1.0,
            "space_before": pf.space_before.pt if pf.space_before else 0,
            "space_after": pf.space_after.pt if pf.space_after else 0,
            "first_line_indent": pf.first_line_indent.cm if pf.first_line_indent else 0,
            "left_indent": pf.left_indent.cm if pf.left_indent else 0,
            "right_indent": pf.right_indent.cm if pf.right_indent else 0,
        }

    # --- Get section (page) setup ---
    section = doc.sections[0]
    page_width_cm = section.page_width.cm
    page_height_cm = section.page_height.cm
    margin_top_cm = section.top_margin.cm
    margin_bottom_cm = section.bottom_margin.cm
    margin_left_cm = section.left_margin.cm
    margin_right_cm = section.right_margin.cm

    # --- Styles ---
    normal = get_paragraph_style("Normal")
    heading1 = get_paragraph_style("Heading 1")
    heading2 = get_paragraph_style("Heading 2")
    
    # Try to detect actual paragraph alignment if style alignment is missing
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and heading1["alignment"] == "left" and p.alignment is not None:
            heading1["alignment"] = ALIGNMENT_NAMES.get(p.alignment, "left")
        elif p.style.name == "Heading 2" and heading2["alignment"] == "left" and p.alignment is not None:
            heading2["alignment"] = ALIGNMENT_NAMES.get(p.alignment, "left")
    
        # --- Page numbering alignment (if present) ---
        page_number_alignment = None
        for p in section.footer.paragraphs:
            for r in p.runs:
                if "PAGE" in r.text.upper():
                    page_number_alignment = ALIGNMENT_NAMES.get(p.alignment, "center")

    # --- Print out ready-to-paste block ---
    print('STYLE_EXTRACTED = {')
    print('    # --- Page setup ---')
    print(f'    "page_width_cm": {page_width_cm:.1f},')
    print(f'    "page_height_cm": {page_height_cm:.1f},')
    print(f'    "margin_top_cm": {margin_top_cm:.1f},')
    print(f'    "margin_bottom_cm": {margin_bottom_cm:.1f},')
    print(f'    "margin_left_cm": {margin_left_cm:.1f},')
    print(f'    "margin_right_cm": {margin_right_cm:.1f},\n')

    print('    # --- Normal (body) style ---')
    print(f'    "body_font_name": "{normal["font_name"]}",')
    print(f'    "body_font_size": {normal["font_size"] or 11},')
    print(f'    "line_spacing": {normal["line_spacing"]},')
    print(f'    "before_space": {normal["space_before"]},')
    print(f'    "after_space": {normal["space_after"]},')
    print(f'    "alignment": "{normal["alignment"]}",')
    print(f'    "first_line_indent_cm": {normal["first_line_indent"]},')
    print(f'    "left_indent_cm": {normal["left_indent"]},')
    print(f'    "right_indent_cm": {normal["right_indent"]},\n')

    print('    # --- Heading 1 ---')
    print(f'    "heading1_font_name": "{heading1["font_name"]}",')
    print(f'    "heading1_font_size": {heading1["font_size"] or normal["font_size"]},')
    print(f'    "heading1_alignment": "{heading1["alignment"]}",')
    print(f'    "heading1_first_line_indent_cm": {heading1["first_line_indent"]},')
    print(f'    "heading1_left_indent_cm": {heading1["left_indent"]},')
    print(f'    "heading1_right_indent_cm": {heading1["right_indent"]},')
    print(f'    "heading1_before_space": {heading1["space_before"]},')
    print(f'    "heading1_after_space": {heading1["space_after"]},\n')

    print('    # --- Heading 2 ---')
    print(f'    "heading2_font_name": "{heading2["font_name"]}",')
    print(f'    "heading2_font_size": {heading2["font_size"] or normal["font_size"]},')
    print(f'    "heading2_alignment": "{heading2["alignment"]}",')
    print(f'    "heading2_first_line_indent_cm": {heading2["first_line_indent"]},')
    print(f'    "heading2_left_indent_cm": {heading2["left_indent"]},')
    print(f'    "heading2_right_indent_cm": {heading2["right_indent"]},')
    print(f'    "heading2_before_space": {heading2["space_before"]},')
    print(f'    "heading2_after_space": {heading2["space_after"]},\n')

    print('    # --- Page number ---')
    print(f'    "page_number_alignment": "{page_number_alignment or "center"}"')
    print('}')



# === Example usage ===
if __name__ == "__main__":
    extract_global_style("il_sapore_del_tempo_2.0.docx")
