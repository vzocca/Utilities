from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH


# === GLOBAL STYLE PROFILES ===
STYLE_A = {
    # --- Page setup ---
    "page_width_cm": 21.6,
    "page_height_cm": 27.9,
    "margin_top_cm": 2.5,
    "margin_bottom_cm": 2.5,
    "margin_left_cm": 2.5,
    "margin_right_cm": 2.5,

    # --- Normal (body) style ---
    "body_font_name": "Times New Roman",
    "body_font_size": 12.0,
    "line_spacing": 2.0,
    "before_space": 0,
    "after_space": 6.0,
    "alignment": "justify",
    "first_line_indent_cm": 0.5,
    "left_indent_cm": 0,
    "right_indent_cm": 0,
    
    # --- Title ---
    "title_font_name": "Times New Roman",
    "title_font_size": 18.0,
    "title_alignment": "center",
    "title_first_line_indent_cm": 0,
    "title_left_indent_cm": 0,
    "title_right_indent_cm": 0,
    "title_before_space": 36.0,
    "title_after_space": 24.0,

    # --- Heading 1 ---
    "heading1_font_name": "Times New Roman",
    "heading1_font_size": 16.0,
    "heading1_alignment": "center",
    "heading1_first_line_indent_cm": 0,
    "heading1_left_indent_cm": 0,
    "heading1_right_indent_cm": 0,
    "heading1_before_space": 24.0,
    "heading1_after_space": 12.0,

    # --- Heading 2 ---
    "heading2_font_name": "Times New Roman",
    "heading2_font_size": 14.0,
    "heading2_alignment": "center",
    "heading2_first_line_indent_cm": 0,
    "heading2_left_indent_cm": 0,
    "heading2_right_indent_cm": 0,
    "heading2_before_space": 18.0,
    "heading2_after_space": 0,
    
    
    # --- Quote ---
    "quote_font_name": "Times New Roman",
    "quote_font_style": "italic",
    "quote_font_size": 14.0,
    "quote_alignment": "center",
    "quote_first_line_indent_cm": 0,
    "quote_left_indent_cm": 0,
    "quote_right_indent_cm": 0,
    "quote_before_space": 18.0,
    "quote_after_space": 0,

    # --- Page number ---
    "page_number_alignment": "center"
}

STYLE_B = {
    # --- Page setup ---
    "page_width_cm": 14.0,
    "page_height_cm": 21.6,
    "margin_top_cm": 2.5,
    "margin_bottom_cm": 2.5,
    "margin_left_cm": 2.5,
    "margin_right_cm": 2.0,

    # --- Normal (body) style ---
    "body_font_name": "Garamond",
    "body_font_size": 12.0,
    "line_spacing": 1.25,
    "before_space": 0,
    "after_space": 6.0,
    "alignment": "justify",
    "first_line_indent_cm": 0.5,
    "left_indent_cm": 0,
    "right_indent_cm": 0,
    
    
    # --- Title ---
    "title_font_name": "Garamond",
    "title_font_size": 18.0,
    "title_alignment": "center",
    "title_first_line_indent_cm": 0,
    "title_left_indent_cm": 0,
    "title_right_indent_cm": 0,
    "title_before_space": 36.0,
    "title_after_space": 24.0,

    # --- Heading 1 ---
    "heading1_font_name": "Garamond",
    "heading1_font_size": 16.0,
    "heading1_alignment": "center",
    "heading1_first_line_indent_cm": 0,
    "heading1_left_indent_cm": 0,
    "heading1_right_indent_cm": 0,
    "heading1_before_space": 24.0,
    "heading1_after_space": 12.0,

    # --- Heading 2 ---
    "heading2_font_name": "Garamond",
    "heading2_font_size": 14.0,
    "heading2_alignment": "center",
    "heading2_first_line_indent_cm": 0,
    "heading2_left_indent_cm": 0,
    "heading2_right_indent_cm": 0,
    "heading2_before_space": 18.0,
    "heading2_after_space": 0,
    
    # --- Quote ---
    "quote_font_name": "Garamond",
    "quote_font_style": "italic",
    "quote_font_size": 14.0,
    "quote_alignment": "center",
    "quote_first_line_indent_cm": 0,
    "quote_left_indent_cm": 0,
    "quote_right_indent_cm": 0,
    "quote_before_space": 18.0,
    "quote_after_space": 0,
    

    # --- Page number ---
    "page_number_alignment": "center"
}

# === HELPER FUNCTIONS ===
def _get_alignment(value):
    """Convert alignment string to docx enum."""
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(value.lower(), WD_ALIGN_PARAGRAPH.LEFT)


# === MAIN FUNCTION ===
def convert_docx_advanced(input_file, output_file, **style):
    """
    Converts one .docx file to another with updated layout, margins, fonts,
    alignment, indentation, and page numbering.
    """
    doc = Document(input_file)

    # --- Page setup ---
    for section in doc.sections:
        section.page_width = Cm(style["page_width_cm"])
        section.page_height = Cm(style["page_height_cm"])
        section.top_margin = Cm(style["margin_top_cm"])
        section.bottom_margin = Cm(style["margin_bottom_cm"])
        section.left_margin = Cm(style["margin_left_cm"])
        section.right_margin = Cm(style["margin_right_cm"])

    # --- Paragraph and font changes ---
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name.lower()

        # --- Heading 1 ---
        if style_name.startswith("heading 1"):
            font_name = style["heading1_font_name"]
            font_size = style["heading1_font_size"]
            fmt = paragraph.paragraph_format
            fmt.alignment = _get_alignment(style["heading1_alignment"])
            fmt.first_line_indent = Cm(style["heading1_first_line_indent_cm"])
            fmt.left_indent = Cm(style["heading1_left_indent_cm"])
            fmt.right_indent = Cm(style["heading1_right_indent_cm"])
            fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            fmt.line_spacing = style["line_spacing"]
            fmt.space_before = Pt(style["heading1_before_space"])
            fmt.space_after = Pt(style["heading1_after_space"])

        # --- Heading 2 ---
        elif style_name.startswith("heading 2"):
            font_name = style["heading2_font_name"]
            font_size = style["heading2_font_size"]
            fmt = paragraph.paragraph_format
            fmt.alignment = _get_alignment(style["heading2_alignment"])
            fmt.first_line_indent = Cm(style["heading2_first_line_indent_cm"])
            fmt.left_indent = Cm(style["heading2_left_indent_cm"])
            fmt.right_indent = Cm(style["heading2_right_indent_cm"])
            fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            fmt.line_spacing = style["line_spacing"]
            fmt.space_before = Pt(style["heading2_before_space"])
            fmt.space_after = Pt(style["heading2_after_space"])
            
        # --- Quote ---
        elif style_name.startswith("quote"):
            font_name = style["quote_font_name"]
            font_size = style["quote_font_size"]
            fmt = paragraph.paragraph_format
            fmt.alignment = _get_alignment(style["quote_alignment"])
            fmt.first_line_indent = Cm(style["quote_first_line_indent_cm"])
            fmt.left_indent = Cm(style["quote_left_indent_cm"])
            fmt.right_indent = Cm(style["quote_right_indent_cm"])
            fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            fmt.line_spacing = style["line_spacing"]
            fmt.space_before = Pt(style["quote_before_space"])
            fmt.space_after = Pt(style["quote_after_space"])
            
        # --- Title ---
        elif style_name.startswith("title"):
            font_name = style["title_font_name"]
            font_size = style["title_font_size"]
            fmt = paragraph.paragraph_format
            fmt.alignment = _get_alignment(style["title_alignment"])
            fmt.first_line_indent = Cm(style["title_first_line_indent_cm"])
            fmt.left_indent = Cm(style["title_left_indent_cm"])
            fmt.right_indent = Cm(style["title_right_indent_cm"])
            fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            fmt.line_spacing = style["line_spacing"]
            fmt.space_before = Pt(style["title_before_space"])
            fmt.space_after = Pt(style["title_after_space"])

        # --- Normal paragraphs ---
        else:
            font_name = style["body_font_name"]
            font_size = style["body_font_size"]
            fmt = paragraph.paragraph_format
            fmt.alignment = _get_alignment(style["alignment"])
            fmt.first_line_indent = Cm(style["first_line_indent_cm"])
            fmt.left_indent = Cm(style["left_indent_cm"])
            fmt.right_indent = Cm(style["right_indent_cm"])
            fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            fmt.line_spacing = style["line_spacing"]
            fmt.space_before = Pt(style["before_space"])
            fmt.space_after = Pt(style["after_space"])

        # --- Apply font formatting to runs ---
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

    # --- Page numbering alignment ---
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            for r in p.runs:
                if "PAGE" in r.text.upper():
                    p.alignment = _get_alignment(style.get("page_number_alignment", "center"))

    doc.save(output_file)
    print(f"✅ Converted successfully → {output_file}")


# === EXAMPLE USAGE ===
if __name__ == "__main__":
    BASE = "C:\\Users\\vzocc\\Documents\\GitHub\\"
    
    DIR = "Il-sapore-del-tempo\\"
    FILE = "Il_sapore_del_tempo.docx"
    # Choose which style profile to apply
    #convert_docx_advanced("C:\\Users\\vzocc\\Documents\\GitHub\\Il-sapore-del-tempo\\il_sapore_del_tempo.docx", "C:\\Users\\vzocc\\Documents\\GitHub\\Il-sapore-del-tempo\\versions\\il_sapore_del_tempo_1.25.docx", **STYLE_B)
    convert_docx_advanced(BASE+DIR+FILE, BASE+DIR+"versions\\"+"il_sapore_del_tempo_1.25.docx", **STYLE_B)
                           

    # Choose which style profile to apply
    #convert_docx_advanced("C:\\Users\\vzocc\\Documents\\GitHub\\Il-sapore-del-tempo\\il_sapore_del_tempo.docx", "C:\\Users\\vzocc\\Documents\\GitHub\\Il-sapore-del-tempo\\versions\\il_sapore_del_tempo_2.0.docx", **STYLE_A)
    convert_docx_advanced("C:\\Users\\vzocc\\Documents\\GitHub\\Il-sapore-del-tempo\\il_sapore_del_tempo.docx", "C:\\Users\\vzocc\\Documents\\GitHub\\Il-sapore-del-tempo\\versions\\il_sapore_del_tempo_2.0.docx", **STYLE_A)
        